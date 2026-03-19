#!/usr/bin/env python3
"""
run_tests.py
~~~~~~~~~~~~
Standalone test runner for Phase 1.
Runs all validation rule tests without requiring pytest.
Uses the conftest fixture loader directly and a simple assertion framework.

Usage: python run_tests.py
"""

import sys
import traceback
import copy
from pathlib import Path
from datetime import date, datetime, timedelta

sys.path.insert(0, str(Path(__file__).parent))

from scopeguard.validator import Validator
from scopeguard.models import (
    NetworkAsset, OutOfScopeAsset, PhysicalLocation, Technique, MaintenanceWindow,
    DeliveryMethod, AuthorizationStatus, TechniqueCategory, DocumentStatus,
    EngagementType, RegulatoryBasis, UsbPayloadType, BlackoutDate,
)
from scopeguard.finding import Severity
from tests.conftest import load_fixture


# ─── Test infrastructure ──────────────────────────────────────────────────────

PASSED = 0
FAILED = 0
ERRORS = 0
_results = []


def validate(engagement):
    return Validator(engagement).validate()


def has_rule(findings, rule_id):
    return any(f.rule_id == rule_id for f in findings)


def findings_for(findings, rule_id):
    return [f for f in findings if f.rule_id == rule_id]


def run_test(name, fn, *args):
    global PASSED, FAILED, ERRORS
    try:
        fn(*args)
        PASSED += 1
        _results.append(("PASS", name, None))
    except AssertionError as e:
        FAILED += 1
        _results.append(("FAIL", name, str(e)))
    except Exception as e:
        ERRORS += 1
        _results.append(("ERROR", name, traceback.format_exc()))


# ─── Fixtures ─────────────────────────────────────────────────────────────────

mcb = load_fixture("mcb")
nexus = load_fixture("nexus_bad")


# ─── VAL-001 ──────────────────────────────────────────────────────────────────

def test_val001_valid_no_trigger():
    assert not has_rule(validate(mcb), "VAL-001")

def test_val001_invalid_cidr_detected():
    eng = copy.deepcopy(mcb)
    eng.in_scope_assets[0].cidr_notation = "999.999.999.999/33"
    assert has_rule(validate(eng), "VAL-001")


# ─── VAL-002 ──────────────────────────────────────────────────────────────────

def test_val002_mcb_no_trigger():
    f = validate(mcb)
    assert not has_rule(f, "VAL-002"), \
        f"VAL-002 fired unexpectedly: {[str(x) for x in findings_for(f, 'VAL-002')]}"

def test_val002_nexus_detected():
    assert has_rule(validate(nexus), "VAL-002"), \
        "VAL-002 did not detect /21 vs 255.255.255.0 mismatch"

def test_val002_specific_mismatch():
    eng = copy.deepcopy(mcb)
    eng.in_scope_assets[0].cidr_notation = "10.10.0.0/22"
    eng.in_scope_assets[0].subnet_mask = "255.255.255.0"
    findings = validate(eng)
    assert has_rule(findings, "VAL-002")
    blocker = next(f for f in findings if f.rule_id == "VAL-002")
    assert blocker.severity == Severity.BLOCK


# ─── VAL-003 ──────────────────────────────────────────────────────────────────

def test_val003_mcb_no_trigger():
    assert not has_rule(validate(mcb), "VAL-003")

def test_val003_nexus_detected():
    assert has_rule(validate(nexus), "VAL-003"), \
        "VAL-003 did not detect range 10-12 stated as 5"

def test_val003_off_by_one():
    eng = copy.deepcopy(mcb)
    asset = next(a for a in eng.in_scope_assets if "Core Banking" in a.asset_name)
    asset.vlan_count_stated = 3   # range 10-13 = 4, not 3
    assert has_rule(validate(eng), "VAL-003")

def test_val003_none_stated_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.in_scope_assets[0].vlan_count_stated = None
    assert not has_rule(validate(eng), "VAL-003")


# ─── VAL-004 ──────────────────────────────────────────────────────────────────

def test_val004_valid_no_trigger():
    assert not has_rule(validate(mcb), "VAL-004")

def test_val004_end_before_start():
    eng = copy.deepcopy(mcb)
    eng.period.authorized_end_date = eng.period.authorized_start_date - timedelta(days=1)
    assert has_rule(validate(eng), "VAL-004")


# ─── VAL-005 ──────────────────────────────────────────────────────────────────

def test_val005_valid_no_trigger():
    assert not has_rule(validate(mcb), "VAL-005")

def test_val005_end_before_start():
    eng = copy.deepcopy(mcb)
    eng.period.active_testing_hours_end = "05:00 ET"
    eng.period.active_testing_hours_start = "22:00 ET"
    assert has_rule(validate(eng), "VAL-005")


# ─── VAL-006 ──────────────────────────────────────────────────────────────────

def test_val006_valid_no_trigger():
    assert not has_rule(validate(mcb), "VAL-006")

def test_val006_final_before_draft():
    eng = copy.deepcopy(mcb)
    eng.period.report_final_due = date(2026, 4, 30)
    eng.period.report_draft_due = date(2026, 5, 2)
    assert has_rule(validate(eng), "VAL-006")


# ─── VAL-007 ──────────────────────────────────────────────────────────────────

def test_val007_valid_ips_no_trigger():
    assert not has_rule(validate(mcb), "VAL-007")

def test_val007_invalid_ip_detected():
    assert has_rule(validate(nexus), "VAL-007")

def test_val007_hostname_rejected():
    eng = copy.deepcopy(mcb)
    eng.contacts[7].authorized_source_ips = ["attacker.example.com"]
    assert has_rule(validate(eng), "VAL-007")


# ─── VAL-008 ──────────────────────────────────────────────────────────────────

def test_val008_valid_emails_no_trigger():
    assert not has_rule(validate(mcb), "VAL-008")

def test_val008_invalid_email_detected():
    assert has_rule(validate(nexus), "VAL-008")

def test_val008_missing_at_sign():
    eng = copy.deepcopy(mcb)
    eng.contacts[0].email = "noatsignexample.com"
    assert has_rule(validate(eng), "VAL-008")


# ─── VAL-009 ──────────────────────────────────────────────────────────────────

def test_val009_valid_conditionals_no_trigger():
    f = validate(mcb)
    assert not has_rule(f, "VAL-009"), \
        f"VAL-009 fired unexpectedly: {[str(x) for x in findings_for(f, 'VAL-009')]}"

def test_val009_vague_condition_detected():
    assert has_rule(validate(nexus), "VAL-009")

def test_val009_missing_workflow():
    eng = copy.deepcopy(mcb)
    t = next(t for t in eng.techniques if t.authorization_status == AuthorizationStatus.CONDITIONAL)
    t.approval_workflow = None
    assert has_rule(validate(eng), "VAL-009")

def test_val009_vague_phrases():
    phrases = ["only if necessary", "if absolutely necessary", "at discretion", "as appropriate"]
    for phrase in phrases:
        eng = copy.deepcopy(mcb)
        t = next(t for t in eng.techniques if t.authorization_status == AuthorizationStatus.CONDITIONAL)
        t.conditions = f"Allowed {phrase}."
        assert has_rule(validate(eng), "VAL-009"), f"Phrase '{phrase}' not detected"


# ─── VAL-010 ──────────────────────────────────────────────────────────────────

def test_val010_valid_no_trigger():
    assert not has_rule(validate(mcb), "VAL-010")

def test_val010_nexus_detected():
    # Nexus has maintenance_window_ref='MW-NONEXISTENT' (ref is set, but window doesn't exist)
    # That triggers XRF-004, not VAL-010.
    # VAL-010 fires when maintenance_window_required=True AND maintenance_window_ref is None.
    findings = validate(nexus)
    assert has_rule(findings, "XRF-004"), \
        "XRF-004 (dangling window ref) should fire for Nexus"

def test_val010_window_required_no_ref():
    eng = copy.deepcopy(mcb)
    eng.techniques.append(Technique(
        technique_id="TECH-DOS-999",
        category=TechniqueCategory.DOS,
        technique_name="Orphaned DoS technique",
        authorization_status=AuthorizationStatus.CONDITIONAL,
        conditions="Only during a defined maintenance window with written client approval.",
        approval_workflow="Client Technical Contact confirms in writing 4 hours before.",
        maintenance_window_required=True,
        maintenance_window_ref=None,
        notification_required=False,
        prohibited=False,
    ))
    assert has_rule(validate(eng), "VAL-010")


# ─── VAL-011 ──────────────────────────────────────────────────────────────────

def test_val011_4_hours_no_trigger():
    assert not has_rule(validate(mcb), "VAL-011")

def test_val011_8_hours_detected():
    assert has_rule(validate(nexus), "VAL-011")

def test_val011_5_hours_detected():
    eng = copy.deepcopy(mcb)
    eng.data_governance.credential_reporting_window_hours = 5
    assert has_rule(validate(eng), "VAL-011")

def test_val011_exactly_4_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.data_governance.credential_reporting_window_hours = 4
    assert not has_rule(validate(eng), "VAL-011")


# ─── VAL-012 ──────────────────────────────────────────────────────────────────

def test_val012_true_no_trigger():
    assert not has_rule(validate(mcb), "VAL-012")

def test_val012_false_detected():
    assert has_rule(validate(nexus), "VAL-012")
    b = next(f for f in validate(nexus) if f.rule_id == "VAL-012")
    assert b.severity == Severity.BLOCK

def test_val012_setting_false():
    eng = copy.deepcopy(mcb)
    eng.data_governance.third_party_disclosure_prohibited = False
    assert has_rule(validate(eng), "VAL-012")


# ─── VAL-013 ──────────────────────────────────────────────────────────────────

def test_val013_inert_no_trigger():
    assert not has_rule(validate(mcb), "VAL-013")

def test_val013_executable_no_auth_detected():
    eng = copy.deepcopy(mcb)
    eng.social_engineering.usb_payload_type = UsbPayloadType.EXECUTABLE
    eng.social_engineering.usb_executable_authorization = None
    assert has_rule(validate(eng), "VAL-013")

def test_val013_executable_with_auth_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.social_engineering.usb_payload_type = UsbPayloadType.EXECUTABLE
    eng.social_engineering.usb_executable_authorization = "Signed by CISO 2026-04-01"
    assert not has_rule(validate(eng), "VAL-013")


# ─── VAL-014 ──────────────────────────────────────────────────────────────────

def test_val014_pending_no_trigger():
    assert not has_rule(validate(mcb), "VAL-014")

def test_val014_nexus_executed_no_sigs_detected():
    assert has_rule(validate(nexus), "VAL-014")

def test_val014_fully_signed_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.identity.document_status = DocumentStatus.EXECUTED
    eng.identity.client_signatory_name = "Sandra K. Whitfield"
    eng.identity.client_signatory_date = date(2026, 4, 1)
    eng.identity.tester_lead_signatory_name = "Marcus T. Holloway"
    eng.identity.tester_lead_signatory_date = date(2026, 4, 1)
    eng.identity.tester_principal_signatory_name = "RSG Director"
    eng.identity.tester_principal_signatory_date = date(2026, 4, 1)
    assert not has_rule(validate(eng), "VAL-014")


# ─── VAL-015 ──────────────────────────────────────────────────────────────────

def test_val015_all_roles_present_no_trigger():
    assert not has_rule(validate(mcb), "VAL-015")

def test_val015_missing_engagement_lead():
    eng = copy.deepcopy(mcb)
    eng.contacts = [c for c in eng.contacts if c.role != "engagement_lead"]
    found = findings_for(validate(eng), "VAL-015")
    assert any("engagement_lead" in f.field_path for f in found)

def test_val015_nexus_multiple_missing():
    found = findings_for(validate(nexus), "VAL-015")
    assert len(found) >= 2


# ─── VAL-016 ──────────────────────────────────────────────────────────────────

def test_val016_mcb_provisioned_flagged():
    found = findings_for(validate(mcb), "VAL-016")
    assert len(found) == 3, f"Expected 3 VAL-016 findings, got {len(found)}"

def test_val016_confirmed_no_trigger():
    eng = copy.deepcopy(mcb)
    for asset in eng.in_scope_assets:
        if asset.delivery_method == DeliveryMethod.CLIENT_PROVISIONED:
            asset.delivery_confirmed = True
            asset.delivery_confirmed_date = date(2026, 4, 4)
            asset.confirmed_address = "203.0.113.1"
    assert not has_rule(validate(eng), "VAL-016")


# ─── VAL-017 ──────────────────────────────────────────────────────────────────

def test_val017_locations_defined_no_trigger():
    assert not has_rule(validate(mcb), "VAL-017")

def test_val017_full_scope_no_locations_detected():
    eng = copy.deepcopy(mcb)
    eng.physical_locations = []
    assert has_rule(validate(eng), "VAL-017")


# ─── VAL-018 ──────────────────────────────────────────────────────────────────

def test_val018_explicit_list_no_trigger():
    assert not has_rule(validate(mcb), "VAL-018")

def test_val018_none_detected():
    eng = copy.deepcopy(mcb)
    eng.social_engineering.excluded_se_targets = None
    assert has_rule(validate(eng), "VAL-018")

def test_val018_empty_list_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.social_engineering.excluded_se_targets = []
    assert not has_rule(validate(eng), "VAL-018")


# ─── VAL-019 ──────────────────────────────────────────────────────────────────

def test_val019_date_defined_no_trigger():
    assert not has_rule(validate(mcb), "VAL-019")

def test_val019_no_delivery_date_detected():
    eng = copy.deepcopy(mcb)
    eng.social_engineering.phishing_target_list_due_date = None
    assert has_rule(validate(eng), "VAL-019")


# ─── VAL-020 ──────────────────────────────────────────────────────────────────

def test_val020_activities_defined_no_trigger():
    assert not has_rule(validate(mcb), "VAL-020")

def test_val020_empty_activities_detected():
    eng = copy.deepcopy(mcb)
    eng.maintenance_windows[0].authorized_activity_refs = []
    assert has_rule(validate(eng), "VAL-020")


# ─── XRF-001 ──────────────────────────────────────────────────────────────────

def test_xrf001_no_overlap_no_trigger():
    assert not has_rule(validate(mcb), "XRF-001")

def test_xrf001_same_cidr_both_lists():
    eng = copy.deepcopy(mcb)
    eng.out_of_scope_assets.append(OutOfScopeAsset(
        asset_name="Duplicate Core Banking",
        cidr_notation="10.10.0.0/22",
        subnet_mask="255.255.252.0",
        description="Same CIDR as Core Banking",
        delivery_method=DeliveryMethod.NETWORK_DISCOVERABLE,
        exclusion_reason="Test",
        third_party_operated=False,
        regulatory_exclusion=False,
    ))
    assert has_rule(validate(eng), "XRF-001")


# ─── XRF-002 ──────────────────────────────────────────────────────────────────

def test_xrf002_no_subnet_no_trigger():
    assert not has_rule(validate(mcb), "XRF-002")

def test_xrf002_in_scope_subnet_of_out_scope():
    eng = copy.deepcopy(mcb)
    eng.out_of_scope_assets.append(OutOfScopeAsset(
        asset_name="Supernet Exclusion",
        cidr_notation="10.10.0.0/16",
        subnet_mask="255.255.0.0",
        description="Contains in-scope Core Banking /22",
        delivery_method=DeliveryMethod.NETWORK_DISCOVERABLE,
        exclusion_reason="Test",
        third_party_operated=False,
        regulatory_exclusion=False,
    ))
    assert has_rule(validate(eng), "XRF-002")


# ─── XRF-003 ──────────────────────────────────────────────────────────────────

def test_xrf003_no_carve_no_trigger():
    assert not has_rule(validate(mcb), "XRF-003")

def test_xrf003_out_scope_subnet_of_in_scope():
    eng = copy.deepcopy(mcb)
    # 10.10.1.0/24 IS a subnet of Core Banking 10.10.0.0/22 (covers 10.10.0.0-10.10.3.255)
    eng.out_of_scope_assets.append(OutOfScopeAsset(
        asset_name="Internal Carve-Out",
        cidr_notation="10.10.1.0/24",
        subnet_mask="255.255.255.0",
        description="Subnet carved out of Core Banking 10.10.0.0/22",
        delivery_method=DeliveryMethod.NETWORK_DISCOVERABLE,
        exclusion_reason="Specific host group exclusion",
        third_party_operated=False,
        regulatory_exclusion=False,
    ))
    assert has_rule(validate(eng), "XRF-003")


# ─── XRF-004 ──────────────────────────────────────────────────────────────────

def test_xrf004_valid_refs_no_trigger():
    assert not has_rule(validate(mcb), "XRF-004")

def test_xrf004_nexus_dangling_ref():
    assert has_rule(validate(nexus), "XRF-004")

def test_xrf004_bad_window_ref():
    eng = copy.deepcopy(mcb)
    t = next(t for t in eng.techniques if t.maintenance_window_ref)
    t.maintenance_window_ref = "MW-GHOST"
    assert has_rule(validate(eng), "XRF-004")


# ─── XRF-005 ──────────────────────────────────────────────────────────────────

def test_xrf005_within_period_no_trigger():
    # MW-003 is in the retest window (May 6) which is covered when retest_included=True
    assert not has_rule(validate(mcb), "XRF-005")

def test_xrf005_window_before_start():
    eng = copy.deepcopy(mcb)
    eng.maintenance_windows[0].date = date(2026, 3, 1)
    assert has_rule(validate(eng), "XRF-005")

def test_xrf005_window_after_end():
    eng = copy.deepcopy(mcb)
    eng.maintenance_windows[0].date = date(2026, 6, 1)
    assert has_rule(validate(eng), "XRF-005")


# ─── XRF-006 ──────────────────────────────────────────────────────────────────

def test_xrf006_valid_refs_no_trigger():
    assert not has_rule(validate(mcb), "XRF-006")

def test_xrf006_dangling_recipient():
    eng = copy.deepcopy(mcb)
    t = next(t for t in eng.techniques if t.notification_required)
    t.notification_recipient_ref = "ghost_role"
    f = findings_for(validate(eng), "XRF-006")
    assert f
    assert f[0].severity == Severity.BLOCK


# ─── XRF-007 ──────────────────────────────────────────────────────────────────

def test_xrf007_valid_retest_no_trigger():
    assert not has_rule(validate(mcb), "XRF-007")

def test_xrf007_retest_before_end():
    eng = copy.deepcopy(mcb)
    eng.period.retest_window_start = date(2026, 4, 20)  # before April 25
    assert has_rule(validate(eng), "XRF-007")


# ─── XRF-008 ──────────────────────────────────────────────────────────────────

def test_xrf008_valid_blackout_no_trigger():
    assert not has_rule(validate(mcb), "XRF-008")

def test_xrf008_blackout_outside_period():
    eng = copy.deepcopy(mcb)
    eng.period.blackout_dates.append(BlackoutDate(date=date(2026, 3, 1), reason="Before start"))
    assert has_rule(validate(eng), "XRF-008")


# ─── XRF-009 ──────────────────────────────────────────────────────────────────

def test_xrf009_manager_defined_no_trigger():
    assert not has_rule(validate(mcb), "XRF-009")

def test_xrf009_missing_manager_clarify():
    eng = copy.deepcopy(mcb)
    eng.contacts = [c for c in eng.contacts if c.role != "physical_security_manager"]
    f = findings_for(validate(eng), "XRF-009")
    assert f
    assert f[0].severity == Severity.CLARIFY


# ─── XRF-010 ──────────────────────────────────────────────────────────────────

def test_xrf010_counsel_defined_no_trigger():
    assert not has_rule(validate(mcb), "XRF-010")

def test_xrf010_missing_counsel_clarify():
    eng = copy.deepcopy(mcb)
    eng.contacts = [c for c in eng.contacts if c.role != "legal_counsel"]
    f = findings_for(validate(eng), "XRF-010")
    assert f
    assert f[0].severity == Severity.CLARIFY


# ─── XRF-011 ──────────────────────────────────────────────────────────────────

def test_xrf011_mcb_cyrusone_unnotified_fires():
    """MCB CyrusOne has facility_notified=False — should trigger."""
    assert has_rule(validate(mcb), "XRF-011")

def test_xrf011_notified_true_no_trigger():
    eng = copy.deepcopy(mcb)
    for loc in eng.physical_locations:
        if loc.facility_third_party:
            loc.facility_notified = True
    assert not has_rule(validate(eng), "XRF-011")


# ─── XRF-012 ──────────────────────────────────────────────────────────────────

def test_xrf012_lead_times_defined_no_trigger():
    assert not has_rule(validate(mcb), "XRF-012")

def test_xrf012_missing_lead_time_clarify():
    eng = copy.deepcopy(mcb)
    t = next(t for t in eng.techniques if t.notification_required)
    t.notification_lead_time_hours = None
    f = findings_for(validate(eng), "XRF-012")
    assert f
    assert f[0].severity == Severity.CLARIFY


# ─── XRF-013 ──────────────────────────────────────────────────────────────────

def test_xrf013_same_16_supernet_noted():
    eng = copy.deepcopy(mcb)
    eng.out_of_scope_assets.append(OutOfScopeAsset(
        asset_name="Same /16",
        cidr_notation="10.10.200.0/24",
        subnet_mask="255.255.255.0",
        description="On same /16 as Core Banking",
        delivery_method=DeliveryMethod.NETWORK_DISCOVERABLE,
        exclusion_reason="Test",
        third_party_operated=False,
        regulatory_exclusion=False,
    ))
    f = findings_for(validate(eng), "XRF-013")
    assert f
    assert all(x.severity == Severity.NOTE for x in f)


# ─── XRF-014 ──────────────────────────────────────────────────────────────────

def test_xrf014_ips_defined_no_trigger():
    assert not has_rule(validate(mcb), "XRF-014")

def test_xrf014_no_ips_noted():
    eng = copy.deepcopy(mcb)
    for c in eng.contacts:
        c.authorized_source_ips = []
    f = findings_for(validate(eng), "XRF-014")
    assert f
    assert f[0].severity == Severity.NOTE


# ─── XRF-015 ──────────────────────────────────────────────────────────────────

def test_xrf015_full_scope_no_hipaa_noted():
    assert has_rule(validate(mcb), "XRF-015")
    f = findings_for(validate(mcb), "XRF-015")
    assert f[0].severity == Severity.NOTE

def test_xrf015_hipaa_listed_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.identity.regulatory_basis.append(RegulatoryBasis.HIPAA)
    assert not has_rule(validate(eng), "XRF-015")


# ─── XRF-016 ──────────────────────────────────────────────────────────────────

def test_xrf016_mcb_has_cde_no_trigger():
    # MCB has 'Card Processing (FIS)' in out-of-scope — keyword 'card' present
    assert not has_rule(validate(mcb), "XRF-016")

def test_xrf016_pci_no_cde_assets_noted():
    eng = copy.deepcopy(mcb)
    eng.out_of_scope_assets = [
        a for a in eng.out_of_scope_assets
        if "card" not in a.asset_name.lower() and "payment" not in a.asset_name.lower()
    ]
    f = findings_for(validate(eng), "XRF-016")
    assert f
    assert f[0].severity == Severity.NOTE

def test_xrf016_no_pci_no_trigger():
    eng = copy.deepcopy(mcb)
    eng.identity.regulatory_basis = [r for r in eng.identity.regulatory_basis
                                      if r.value != "PCI-DSS"]
    assert not has_rule(validate(eng), "XRF-016")


# ─── Phase 1 Milestones ────────────────────────────────────────────────────────

def test_milestone_mcb_zero_blockers():
    """MCB fixture must produce zero BLOCK findings."""
    findings = validate(mcb)
    blockers = findings.blockers()
    assert not blockers, \
        f"MCB produced {len(blockers)} unexpected BLOCK(s):\n" + \
        "\n".join(f"  {b}" for b in blockers)

def test_milestone_nexus_spec_errors_all_detected():
    """Phase 1 milestone — all spec-documented Nexus errors detected.
    Nexus has maintenance_window_ref='MW-NONEXISTENT' (set but points to nonexistent window),
    so XRF-004 fires rather than VAL-010. Both indicate the same problem."""
    findings = validate(nexus)
    assert has_rule(findings, "VAL-002"), "CIDR/mask mismatch not detected"
    assert has_rule(findings, "VAL-003"), "VLAN count mismatch not detected"
    assert has_rule(findings, "VAL-014"), "Executed without signatures not detected"
    assert has_rule(findings, "XRF-004") or has_rule(findings, "VAL-010"), \
        "Maintenance window problem not detected"

def test_milestone_finding_list_summary():
    findings = validate(mcb)
    counts = findings.count()
    # MCB should have zero BLOCKs, some MISSING (unconfirmed assets), some NOTEs
    assert counts["BLOCK"] == 0
    assert counts["MISSING"] >= 3  # 3 unconfirmed client-provisioned assets


# ─── Runner ───────────────────────────────────────────────────────────────────



# ─── V2: Form Builder field types ─────────────────────────────────────────────

def test_fb_engagement_type_buttongroup():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("identity")}
    assert fields["engagement_type"]["type"] == "buttongroup"
    assert len(fields["engagement_type"]["options"]) == 6

def test_fb_classification_buttongroup():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("identity")}
    assert fields["classification"]["type"] == "buttongroup"
    assert "confidential" in fields["classification"]["options"]

def test_fb_device_type_multiselect_in_scope():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("in_scope_assets")}
    assert fields["device_type"]["type"] == "multiselect"
    assert len(fields["device_type"]["options"]) >= 20

def test_fb_device_type_multiselect_out_scope():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("out_of_scope_assets")}
    assert fields["device_type"]["type"] == "multiselect"

def test_fb_authorized_activities_multiselect():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("physical_locations")}
    assert fields["authorized_activities"]["type"] == "multiselect"
    assert len(fields["authorized_activities"]["options"]) == 17

def test_fb_all_multiselect_have_options():
    from app.form_builder import get_section_fields
    for section in ["identity", "period", "in_scope_assets",
                    "out_of_scope_assets", "physical_locations"]:
        for f in get_section_fields(section):
            if f["type"] == "multiselect":
                assert f.get("options"), f"[{section}] {f['name']} has no options"

def test_fb_client_asset_acknowledged_field():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("in_scope_assets")}
    assert "client_asset_list_acknowledged" in fields
    assert fields["client_asset_list_acknowledged"]["type"] == "checkbox"

def test_fb_techniques_returns_empty():
    from app.form_builder import get_section_fields
    assert get_section_fields("techniques") == []

# ─── V2: Technique catalog schema ─────────────────────────────────────────────

def test_technique_schema_has_catalog():
    import yaml
    from pathlib import Path
    schema = yaml.safe_load(
        (Path("schema") / "techniques.yaml").read_text()
    )
    assert "catalog" in schema

def test_technique_schema_seven_categories():
    import yaml
    from pathlib import Path
    schema = yaml.safe_load((Path("schema") / "techniques.yaml").read_text())
    assert len(schema["catalog"]) == 7

def test_technique_ids_unique():
    import yaml
    from pathlib import Path
    schema = yaml.safe_load((Path("schema") / "techniques.yaml").read_text())
    ids = [t["id"] for cat in schema["catalog"].values() for t in cat["techniques"]]
    assert len(ids) == len(set(ids)), "Duplicate technique IDs"

def test_technique_count_at_least_40():
    import yaml
    from pathlib import Path
    schema = yaml.safe_load((Path("schema") / "techniques.yaml").read_text())
    total = sum(len(c["techniques"]) for c in schema["catalog"].values())
    assert total >= 40, f"Expected >= 40 techniques, got {total}"

def test_each_technique_has_required_fields():
    import yaml
    from pathlib import Path
    schema = yaml.safe_load((Path("schema") / "techniques.yaml").read_text())
    for cat_key, cat in schema["catalog"].items():
        for tech in cat["techniques"]:
            for field in ("id", "name", "mitre", "nist", "ptes"):
                assert field in tech, f"{cat_key} technique missing {field!r}: {tech}"

# ─── V2: Physical location activities ─────────────────────────────────────────

def test_physical_activities_snake_case():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("physical_locations")}
    for opt in fields["authorized_activities"]["options"]:
        assert " " not in opt, f"Not snake_case: {opt!r}"
        assert opt == opt.lower(), f"Not lowercase: {opt!r}"

def test_physical_activities_known_items():
    from app.form_builder import get_section_fields
    fields = {f["name"]: f for f in get_section_fields("physical_locations")}
    opts = fields["authorized_activities"]["options"]
    for expected in ["tailgating_entry", "badge_cloning", "server_room_access",
                     "soc_access_attempt", "dumpster_reconnaissance"]:
        assert expected in opts, f"Missing activity: {expected}"

# ─── V2: Document generation ───────────────────────────────────────────────────

def _get_doc_texts():
    import io, json
    from pathlib import Path
    from app.hydrator import hydrate
    from app.generator import generate_sow, generate_roe
    from docx import Document
    data = json.loads((Path("tests/fixtures/mcb.json")).read_text())
    sections = ["identity","period","contacts","in_scope_assets","out_of_scope_assets",
                "physical_locations","techniques","maintenance_windows",
                "data_governance","social_engineering"]
    eng = hydrate({s: data[s] for s in sections})
    sow = generate_sow(eng)
    roe = generate_roe(eng)
    def doc_text(b):
        d = Document(io.BytesIO(b))
        return (" ".join(p.text for p in d.paragraphs) +
                " ".join(c.text for t in d.tables for r in t.rows
                         for c in r.cells for p in c.paragraphs))
    return doc_text(sow), doc_text(roe), sow, roe

def test_gen_sow_non_empty():
    _, _, sow, _ = _get_doc_texts()
    assert len(sow) > 10_000

def test_gen_roe_non_empty():
    _, _, _, roe = _get_doc_texts()
    assert len(roe) > 10_000

def test_gen_sow_has_device_type_col():
    sow_text, _, _, _ = _get_doc_texts()
    assert "Type" in sow_text

def test_gen_sow_has_disclaimer():
    sow_text, _, _, _ = _get_doc_texts()
    assert "UNDISCLOSED DEVICE DISCLAIMER" in sow_text

def test_gen_sow_has_cde_section():
    sow_text, _, _, _ = _get_doc_texts()
    assert "Cardholder Data Environment" in sow_text

def test_gen_sow_has_physical_activities():
    sow_text, _, _, _ = _get_doc_texts()
    assert any(t in sow_text for t in ["Tailgating", "Badge", "Dumpster"])

def test_gen_roe_has_ids_ips():
    _, roe_text, _, _ = _get_doc_texts()
    assert "IDS" in roe_text or "IPS" in roe_text

def test_gen_roe_has_maintenance_windows():
    _, roe_text, _, _ = _get_doc_texts()
    assert "Maintenance Window" in roe_text

def test_gen_sow_has_confidential():
    sow_text, _, _, _ = _get_doc_texts()
    assert "CONFIDENTIAL" in sow_text.upper()

# ─── V2: Route integration ─────────────────────────────────────────────────────

def _setup_test_client():
    import json
    from pathlib import Path
    from app import create_app
    from app.storage import save_section
    data = json.loads((Path("tests/fixtures/mcb.json")).read_text())
    app = create_app()
    client = app.test_client().__enter__()
    resp = client.post("/new", follow_redirects=False)
    eng_id = resp.headers["Location"].split("/")[2]
    for s in ["identity","period","contacts","in_scope_assets","out_of_scope_assets",
              "physical_locations","techniques","maintenance_windows",
              "data_governance","social_engineering"]:
        save_section(eng_id, s, data[s])
    return client, eng_id

def test_route_techniques_renders():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/section/techniques")
    assert r.status_code == 200
    assert b"tech-category" in r.data

def test_route_techniques_has_7_cats():
    import re
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/section/techniques")
    cats = re.findall(rb'class="tech-cat-header"', r.data)
    assert len(cats) == 7

def test_route_identity_has_buttongroup():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/section/identity")
    assert b"btn-group-option" in r.data

def test_route_physical_starts_expanded():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/section/physical_locations")
    assert b"display:block" in r.data

def test_route_in_scope_has_30day():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/section/in_scope_assets")
    assert b"30 days" in r.data or b"client_asset_list_acknowledged" in r.data

def test_route_preflight_renders():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/preflight")
    assert r.status_code == 200

def test_route_sow_generates():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/generate/sow")
    assert r.status_code == 200 and len(r.data) > 10_000

def test_route_roe_generates():
    client, eng_id = _setup_test_client()
    r = client.get(f"/engagement/{eng_id}/generate/roe")
    assert r.status_code == 200 and len(r.data) > 10_000


ALL_TESTS = [
    # VAL rules
    (test_val001_valid_no_trigger, "VAL-001: valid CIDRs no trigger"),
    (test_val001_invalid_cidr_detected, "VAL-001: invalid CIDR detected"),
    (test_val002_mcb_no_trigger, "VAL-002: MCB masks agree no trigger"),
    (test_val002_nexus_detected, "VAL-002: Nexus /21 vs /24 mask mismatch"),
    (test_val002_specific_mismatch, "VAL-002: /22 with /24 mask detected"),
    (test_val003_mcb_no_trigger, "VAL-003: MCB VLAN counts correct"),
    (test_val003_nexus_detected, "VAL-003: Nexus count 5 for range 10-12 detected"),
    (test_val003_off_by_one, "VAL-003: off-by-one count detected"),
    (test_val003_none_stated_no_trigger, "VAL-003: None stated no trigger"),
    (test_val004_valid_no_trigger, "VAL-004: valid dates no trigger"),
    (test_val004_end_before_start, "VAL-004: end before start detected"),
    (test_val005_valid_no_trigger, "VAL-005: valid hours no trigger"),
    (test_val005_end_before_start, "VAL-005: hours end before start detected"),
    (test_val006_valid_no_trigger, "VAL-006: valid report dates no trigger"),
    (test_val006_final_before_draft, "VAL-006: final before draft detected"),
    (test_val007_valid_ips_no_trigger, "VAL-007: valid IPs no trigger"),
    (test_val007_invalid_ip_detected, "VAL-007: invalid IP in Nexus detected"),
    (test_val007_hostname_rejected, "VAL-007: hostname rejected"),
    (test_val008_valid_emails_no_trigger, "VAL-008: valid emails no trigger"),
    (test_val008_invalid_email_detected, "VAL-008: invalid email in Nexus detected"),
    (test_val008_missing_at_sign, "VAL-008: missing @ sign detected"),
    (test_val009_valid_conditionals_no_trigger, "VAL-009: valid conditionals no trigger"),
    (test_val009_vague_condition_detected, "VAL-009: Nexus vague condition detected"),
    (test_val009_missing_workflow, "VAL-009: missing workflow detected"),
    (test_val009_vague_phrases, "VAL-009: vague phrases rejected"),
    (test_val010_valid_no_trigger, "VAL-010: valid window refs no trigger"),
    (test_val010_nexus_detected, "VAL-010: Nexus missing window detected"),
    (test_val010_window_required_no_ref, "VAL-010: window_required=True no ref"),
    (test_val011_4_hours_no_trigger, "VAL-011: 4h window no trigger"),
    (test_val011_8_hours_detected, "VAL-011: 8h window in Nexus detected"),
    (test_val011_5_hours_detected, "VAL-011: 5h window detected"),
    (test_val011_exactly_4_no_trigger, "VAL-011: exactly 4h no trigger"),
    (test_val012_true_no_trigger, "VAL-012: true value no trigger"),
    (test_val012_false_detected, "VAL-012: false in Nexus detected + BLOCK"),
    (test_val012_setting_false, "VAL-012: setting false detected"),
    (test_val013_inert_no_trigger, "VAL-013: inert payload no trigger"),
    (test_val013_executable_no_auth_detected, "VAL-013: executable no auth detected"),
    (test_val013_executable_with_auth_no_trigger, "VAL-013: executable with auth no trigger"),
    (test_val014_pending_no_trigger, "VAL-014: pending status no trigger"),
    (test_val014_nexus_executed_no_sigs_detected, "VAL-014: Nexus executed unsigned detected"),
    (test_val014_fully_signed_no_trigger, "VAL-014: fully signed no trigger"),
    (test_val015_all_roles_present_no_trigger, "VAL-015: all roles present no trigger"),
    (test_val015_missing_engagement_lead, "VAL-015: missing engagement_lead detected"),
    (test_val015_nexus_multiple_missing, "VAL-015: Nexus multiple missing roles"),
    (test_val016_mcb_provisioned_flagged, "VAL-016: 3 unconfirmed provisioned assets"),
    (test_val016_confirmed_no_trigger, "VAL-016: confirmed assets no trigger"),
    (test_val017_locations_defined_no_trigger, "VAL-017: locations defined no trigger"),
    (test_val017_full_scope_no_locations_detected, "VAL-017: full_scope no locations"),
    (test_val018_explicit_list_no_trigger, "VAL-018: explicit list no trigger"),
    (test_val018_none_detected, "VAL-018: None targets detected"),
    (test_val018_empty_list_no_trigger, "VAL-018: empty list no trigger"),
    (test_val019_date_defined_no_trigger, "VAL-019: delivery date defined no trigger"),
    (test_val019_no_delivery_date_detected, "VAL-019: no delivery date detected"),
    (test_val020_activities_defined_no_trigger, "VAL-020: activities defined no trigger"),
    (test_val020_empty_activities_detected, "VAL-020: empty activities detected"),
    # XRF rules
    (test_xrf001_no_overlap_no_trigger, "XRF-001: no CIDR overlap no trigger"),
    (test_xrf001_same_cidr_both_lists, "XRF-001: same CIDR in both lists detected"),
    (test_xrf002_no_subnet_no_trigger, "XRF-002: no subnet relationship no trigger"),
    (test_xrf002_in_scope_subnet_of_out_scope, "XRF-002: in-scope subnet of out-scope"),
    (test_xrf003_no_carve_no_trigger, "XRF-003: no carve-out no trigger"),
    (test_xrf003_out_scope_subnet_of_in_scope, "XRF-003: out-scope subnet of in-scope"),
    (test_xrf004_valid_refs_no_trigger, "XRF-004: valid window refs no trigger"),
    (test_xrf004_nexus_dangling_ref, "XRF-004: Nexus dangling MW ref detected"),
    (test_xrf004_bad_window_ref, "XRF-004: bad window ref detected"),
    (test_xrf005_within_period_no_trigger, "XRF-005: windows within period no trigger"),
    (test_xrf005_window_before_start, "XRF-005: window before start detected"),
    (test_xrf005_window_after_end, "XRF-005: window after end detected"),
    (test_xrf006_valid_refs_no_trigger, "XRF-006: valid notification refs no trigger"),
    (test_xrf006_dangling_recipient, "XRF-006: dangling recipient ref detected"),
    (test_xrf007_valid_retest_no_trigger, "XRF-007: valid retest window no trigger"),
    (test_xrf007_retest_before_end, "XRF-007: retest before end detected"),
    (test_xrf008_valid_blackout_no_trigger, "XRF-008: valid blackout no trigger"),
    (test_xrf008_blackout_outside_period, "XRF-008: blackout outside period detected"),
    (test_xrf009_manager_defined_no_trigger, "XRF-009: manager defined no trigger"),
    (test_xrf009_missing_manager_clarify, "XRF-009: missing manager CLARIFY"),
    (test_xrf010_counsel_defined_no_trigger, "XRF-010: counsel defined no trigger"),
    (test_xrf010_missing_counsel_clarify, "XRF-010: missing counsel CLARIFY"),
    (test_xrf011_mcb_cyrusone_unnotified_fires, "XRF-011: CyrusOne unnotified fires"),
    (test_xrf011_notified_true_no_trigger, "XRF-011: notified=True no trigger"),
    (test_xrf012_lead_times_defined_no_trigger, "XRF-012: lead times defined no trigger"),
    (test_xrf012_missing_lead_time_clarify, "XRF-012: missing lead time CLARIFY"),
    (test_xrf013_same_16_supernet_noted, "XRF-013: same /16 supernet NOTE"),
    (test_xrf014_ips_defined_no_trigger, "XRF-014: IPs defined no trigger"),
    (test_xrf014_no_ips_noted, "XRF-014: no IPs with network techniques NOTE"),
    (test_xrf015_full_scope_no_hipaa_noted, "XRF-015: full_scope no HIPAA NOTE"),
    (test_xrf015_hipaa_listed_no_trigger, "XRF-015: HIPAA listed no trigger"),
    (test_xrf016_mcb_has_cde_no_trigger, "XRF-016: MCB has CDE keywords no trigger"),
    (test_xrf016_pci_no_cde_assets_noted, "XRF-016: PCI-DSS no CDE NOTE"),
    (test_xrf016_no_pci_no_trigger, "XRF-016: no PCI-DSS no trigger"),
    # Milestones
    (test_milestone_mcb_zero_blockers, "MILESTONE: MCB produces zero BLOCK findings"),
    (test_milestone_nexus_spec_errors_all_detected, "MILESTONE: all Nexus spec errors detected"),
    (test_milestone_finding_list_summary, "MILESTONE: FindingList summary counts correct"),
    # V2: Form builder
    (test_fb_engagement_type_buttongroup,    "V2: engagement_type is buttongroup"),
    (test_fb_classification_buttongroup,     "V2: classification is buttongroup"),
    (test_fb_device_type_multiselect_in_scope, "V2: in_scope device_type is multiselect"),
    (test_fb_device_type_multiselect_out_scope,"V2: out_scope device_type is multiselect"),
    (test_fb_authorized_activities_multiselect,"V2: physical activities is multiselect with 17 opts"),
    (test_fb_all_multiselect_have_options,   "V2: all multiselect fields have options"),
    (test_fb_client_asset_acknowledged_field,"V2: client_asset_list_acknowledged field exists"),
    (test_fb_techniques_returns_empty,       "V2: techniques form_builder returns []"),
    # V2: Technique catalog
    (test_technique_schema_has_catalog,      "V2: techniques schema has catalog key"),
    (test_technique_schema_seven_categories, "V2: catalog has 7 categories"),
    (test_technique_ids_unique,              "V2: all technique IDs are unique"),
    (test_technique_count_at_least_40,       "V2: catalog has >= 40 techniques"),
    (test_each_technique_has_required_fields,"V2: every technique has id/name/mitre/nist/ptes"),
    # V2: Physical activities
    (test_physical_activities_snake_case,    "V2: activity options are snake_case"),
    (test_physical_activities_known_items,   "V2: known activity options present"),
    # V2: Document generation
    (test_gen_sow_non_empty,                 "V2: SOW generates non-empty document"),
    (test_gen_roe_non_empty,                 "V2: ROE generates non-empty document"),
    (test_gen_sow_has_device_type_col,       "V2: SOW has device type column"),
    (test_gen_sow_has_disclaimer,            "V2: SOW has undisclosed device disclaimer"),
    (test_gen_sow_has_cde_section,           "V2: SOW has PCI-DSS CDE section"),
    (test_gen_sow_has_physical_activities,   "V2: SOW has physical activity list"),
    (test_gen_roe_has_ids_ips,               "V2: ROE has IDS/IPS status"),
    (test_gen_roe_has_maintenance_windows,   "V2: ROE has maintenance windows section"),
    (test_gen_sow_has_confidential,          "V2: SOW has CONFIDENTIAL classification"),
    # V2: Route integration
    (test_route_techniques_renders,          "V2: techniques section renders 200"),
    (test_route_techniques_has_7_cats,       "V2: technique matrix has 7 categories"),
    (test_route_identity_has_buttongroup,    "V2: identity has buttongroup fields"),
    (test_route_physical_starts_expanded,    "V2: physical locations start expanded"),
    (test_route_in_scope_has_30day,          "V2: in_scope has 30-day notice"),
    (test_route_preflight_renders,           "V2: preflight renders 200"),
    (test_route_sow_generates,               "V2: SOW generation route works"),
    (test_route_roe_generates,               "V2: ROE generation route works"),
]


def main():
    print(f"\nScopeGuard v2 — Complete Test Suite")
    print(f"{'─' * 60}")

    for fn, name in ALL_TESTS:
        run_test(name, fn)

    print()
    for status, name, msg in _results:
        marker = "✓" if status == "PASS" else ("✗" if status == "FAIL" else "!")
        print(f"  {marker} {name}")
        if msg and status != "PASS":
            # Indent the error
            for line in msg.strip().splitlines():
                print(f"      {line}")

    print(f"\n{'─' * 60}")
    print(f"  {PASSED} passed  |  {FAILED} failed  |  {ERRORS} errors")
    print(f"  Total: {len(ALL_TESTS)} tests")
    print()

    return 0 if (FAILED + ERRORS) == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
