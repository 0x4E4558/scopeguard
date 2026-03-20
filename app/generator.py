"""
app.generator
~~~~~~~~~~~~~
Phase 3 — Document generation.
Produces two .docx files from a validated Engagement object:
  - Scope of Work (SOW)
  - Rules of Engagement (ROE)

Uses python-docx. Mirrors the structure of the MCB sample documents.
"""

from __future__ import annotations
import io
from datetime import date, datetime
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

from scopeguard.models import (
    Engagement, AuthorizationStatus,
)
from app.legal import get_sow_legal_clauses, get_roe_legal_clauses


# ── Colour palette ─────────────────────────────────────────────────────────────
DARK   = RGBColor(0x1a, 0x1a, 0x2e)   # near-black for headings
MID    = RGBColor(0x1f, 0x4e, 0x79)   # navy for section headers / table heads
LIGHT  = RGBColor(0xd6, 0xe4, 0xf0)   # pale blue for table header fill
RULE   = RGBColor(0x2e, 0x75, 0xb6)   # blue rule line
GREY   = RGBColor(0xf2, 0xf2, 0xf2)   # alternating row fill
WHITE  = RGBColor(0xff, 0xff, 0xff)
RED_T  = RGBColor(0xC0, 0x00, 0x00)   # CONFIDENTIAL stamp


# ── Header / Footer ───────────────────────────────────────────────────────────

def _add_header_footer(doc_section, classification: str, doc_title: str,
                        eng_id: str, version: str) -> None:
    """Add classification header and page-number footer to a document section."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    # ── Header ──
    header = doc_section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Classification label
    r1 = hp.add_run(f"{'★  ' if classification in ('confidential','restricted') else ''}"
                    f"{classification.upper()} — {doc_title}"
                    f"{'  ★' if classification in ('confidential','restricted') else ''}")
    r1.font.size = Pt(8)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if classification in ('confidential','restricted')                          else RGBColor(0x1F, 0x4E, 0x79)
    r1.font.name = 'Arial'
    # Bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Footer ──
    footer = doc_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Top border
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'), 'single')
    ftop.set(qn('w:sz'), '4')
    ftop.set(qn('w:space'), '1')
    ftop.set(qn('w:color'), '2E75B6')
    fpBdr.append(ftop)
    fpPr.append(fpBdr)

    # Left: engagement ID   Center: page N of M   Right: version
    r_left = fp.add_run(f"{eng_id}")
    r_left.font.size = Pt(8); r_left.font.name = 'Arial'
    r_left.font.color.rgb = RGBColor(0x7E, 0x8A, 0x96)

    # Tab to center
    fp.add_run("	")
    # Page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r_page = fp.add_run()
    r_page.font.size = Pt(8); r_page.font.name = 'Arial'
    r_page.font.color.rgb = RGBColor(0x7E, 0x8A, 0x96)
    r_page._r.append(fldChar1)
    r_page._r.append(instrText)
    r_page._r.append(fldChar2)

    fp.add_run(" of ").font.size = Pt(8)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_total = fp.add_run()
    r_total.font.size = Pt(8); r_total.font.name = 'Arial'
    r_total.font.color.rgb = RGBColor(0x7E, 0x8A, 0x96)
    r_total._r.append(fldChar3)
    r_total._r.append(instrText2)
    r_total._r.append(fldChar4)

    # Tab to right
    fp.add_run("	")
    r_ver = fp.add_run(f"v{version}")
    r_ver.font.size = Pt(8); r_ver.font.name = 'Arial'
    r_ver.font.color.rgb = RGBColor(0x7E, 0x8A, 0x96)

    # Set tab stops: center at 3.25", right at 6.5"
    from docx.oxml import OxmlElement as OE
    tabs = OE('w:tabs')
    t1 = OE('w:tab'); t1.set(qn('w:val'),'center'); t1.set(qn('w:pos'),'4680')
    t2 = OE('w:tab'); t2.set(qn('w:val'),'right');  t2.set(qn('w:pos'),'9360')
    tabs.append(t1); tabs.append(t2)
    fpPr.append(tabs)


# ── Legal clause renderer ─────────────────────────────────────────────────────

def _render_clauses(doc, clauses: list[dict], start_num: int) -> None:
    """Render a list of legal clause dicts into the document."""
    for i, clause in enumerate(clauses, start=start_num):
        _heading(doc, f"{i}. {clause['title']}")
        for para in clause.get('body', []):
            if para.isupper() and len(para) > 20:
                _body(doc, para, bold=True)
            else:
                _body(doc, para)
        for bullet in clause.get('bullets', []):
            _bullet(doc, bullet)
        for para in clause.get('body2', []):
            if para.isupper() and len(para) > 20:
                _body(doc, para, bold=True)
            else:
                _body(doc, para)
        for bullet in clause.get('bullets2', []):
            _bullet(doc, bullet)
        for para in clause.get('body3', []):
            _body(doc, para)


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_border_bottom(para, color: str = "2E75B6", size: int = 12):
    """Add a bottom border rule line to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_col_width(table, col_idx: int, width_inches: float):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── Style helpers ──────────────────────────────────────────────────────────────

def _heading(doc, text: str, level: int = 1):
    """Add a numbered-style heading with a bottom rule."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.name = 'Arial'
    _para_border_bottom(p, color="2E75B6", size=8 if level == 1 else 4)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _subheading(doc, text: str):
    return _heading(doc, text, level=2)


def _body(doc, text: str, bold: bool = False, italic: bool = False, color: Optional[RGBColor] = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def _bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def _notice_box(doc, text: str):
    """A shaded notice paragraph (CONFIDENTIAL notice, warnings)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = MID
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    return p


def _make_table(doc, headers: list[str], col_widths: list[float], data: list[list[str]],
                zebra: bool = True) -> None:
    """
    Add a formatted table.
    headers: column header strings
    col_widths: column widths in inches (must sum to ~7.5 for letter with 0.5" margins each side)
    data: list of rows, each a list of cell strings
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height = Cm(0.7)
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        _set_cell_bg(cell, "1F4E79")
        _set_cell_borders(cell, "1F4E79")
        cell.width = Inches(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = 'Arial'

    # Data rows
    for r_idx, row_data in enumerate(data):
        row = table.add_row()
        fill = "F2F2F2" if (zebra and r_idx % 2 == 1) else "FFFFFF"
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            _set_cell_bg(cell, fill)
            _set_cell_borders(cell, "CCCCCC")
            cell.width = Inches(col_widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val) if val is not None else "")
            run.font.size = Pt(9)
            run.font.name = 'Arial'

    doc.add_paragraph()  # spacing after table


def _signature_table(doc, left_label: str, right_label: str):
    """Two-column signature block."""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    col_w = 3.5

    labels = [
        [left_label,         right_label],
        ["Printed Name / Title", "Printed Name / Title"],
        ["Date",             "Date"],
        ["",                 ""],
    ]
    for r_idx, row_labels in enumerate(labels):
        row = table.rows[r_idx]
        row.height = Cm(1.0 if r_idx > 0 else 0.8)
        for c_idx, lbl in enumerate(row_labels):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, "F2F2F2" if r_idx == 0 else "FFFFFF")
            _set_cell_borders(cell, "CCCCCC")
            cell.width = Inches(col_w)
            p = cell.paragraphs[0]
            run = p.add_run(lbl)
            run.bold = (r_idx == 0)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            if r_idx == 0:
                run.font.color.rgb = MID


def _cover_table(doc, rows: list[tuple[str, str]]):
    """Two-column metadata table for document cover."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for r_idx, (label, value) in enumerate(rows):
        row = table.rows[r_idx]
        lbl_cell = row.cells[0]
        val_cell = row.cells[1]
        _set_cell_bg(lbl_cell, "1F4E79")
        _set_cell_bg(val_cell, "FFFFFF")
        _set_cell_borders(lbl_cell, "1F4E79")
        _set_cell_borders(val_cell, "CCCCCC")
        lbl_cell.width = Inches(2.2)
        val_cell.width = Inches(5.0)
        lp = lbl_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True; lr.font.color.rgb = WHITE
        lr.font.size = Pt(9); lr.font.name = 'Arial'
        vp = val_cell.paragraphs[0]
        vr = vp.add_run(str(value) if value else "")
        vr.font.size = Pt(9); vr.font.name = 'Arial'
    doc.add_paragraph()


def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, datetime):
        return d.strftime("%B %d, %Y")
    if isinstance(d, date):
        return d.strftime("%B %d, %Y")
    s = str(d)
    return s[:10] if len(s) > 10 else s


def _v(val, fallback: str = "[TBD]") -> str:
    """Return val or fallback if empty/None."""
    if val is None or val == "" or val == []:
        return fallback
    return str(val)


def _contact_by_role(eng: Engagement, role: str):
    return eng.contact_by_role(role)


def _contact_display(contact) -> str:
    if not contact:
        return "[Not defined]"
    parts = [contact.full_name, contact.title]
    return ", ".join(p for p in parts if p)


# ══════════════════════════════════════════════════════════════════════════════
# SOW GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def generate_sow(eng: Engagement) -> bytes:
    """Generate the Scope of Work document. Returns bytes."""
    doc = DocxDocument()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    id_  = eng.identity
    per  = eng.period
    dg   = eng.data_governance

    # ── Headers / Footers ──────────────────────────────────────────────────────
    _add_header_footer(
        doc.sections[0],
        classification=id_.classification.value,
        doc_title="SCOPE OF WORK",
        eng_id=id_.engagement_id,
        version=id_.document_version,
    )

    # ── Cover ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PENETRATION TEST")
    r.bold = True; r.font.size = Pt(20); r.font.name = 'Arial'
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCOPE OF WORK & AUTHORIZATION AGREEMENT")
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Arial'
    r.font.color.rgb = MID

    doc.add_paragraph()

    _cover_table(doc, [
        ("Client Organization",  id_.client_org_legal_name),
        ("Engagement ID",        id_.engagement_id),
        ("Engagement Type",      id_.engagement_type.value.replace("_", " ").title()),
        ("Classification",       f"{id_.classification.value.upper()} — RESTRICTED"),
        ("Document Version",     id_.document_version),
        ("Prepared By",          id_.prepared_by),
        ("Prepared Date",        _fmt_date(id_.prepared_date)),
        ("Document Status",      id_.document_status.value.upper().replace("_", " ")),
    ])

    _notice_box(doc,
        "NOTICE: This document contains sensitive security information. "
        "Distribution is strictly limited to signatories and their designated "
        "representatives. Unauthorized disclosure may constitute a violation of "
        "applicable law.")

    _notice_box(doc,
        "STANDARDS ALIGNMENT: This document was produced using ScopeGuard and aligns "
        "with MITRE ATT&CK Enterprise v15, NIST SP 800-115, PTES, OWASP Testing Guide v4.2, "
        "CVSS v3.1, CVE/NVD, and CWE. Technique references in the accompanying Rules of "
        "Engagement use ATT&CK Tactic and Technique IDs. See ROE Appendix A.")

    doc.add_paragraph()

    # ── 1. Engagement Overview ─────────────────────────────────────────────────
    _heading(doc, "1. Engagement Overview")
    _subheading(doc, "1.1 Purpose and Authorization")
    _body(doc,
        f"This Scope of Work and Authorization Agreement (hereinafter 'Agreement') "
        f"constitutes written authorization by {id_.client_org_legal_name} "
        f"(hereinafter 'Client') for {id_.testing_firm_legal_name} "
        f"(hereinafter 'Testing Team') to conduct a {id_.engagement_type.value.replace('_',' ')} "
        f"penetration test against Client-owned and Client-operated information systems "
        f"and physical facilities as defined herein.")

    if id_.regulatory_basis:
        bases = ", ".join(r.value for r in id_.regulatory_basis)
        _body(doc,
            f"This Agreement is entered into pursuant to the Client's obligations under "
            f"{bases} and the Client's internal information security policy mandating "
            f"periodic penetration testing of critical systems.")

    _body(doc,
        "No testing activity may commence prior to full execution of this Agreement "
        "and the accompanying Rules of Engagement document. Commencement of testing "
        "without both documents fully executed by authorized representatives of both "
        "parties constitutes unauthorized access under applicable federal and state "
        "computer fraud statutes.",
        bold=True)

    _subheading(doc, "1.2 Engagement Identification")
    meta_rows = [
        ("Engagement ID",   id_.engagement_id),
        ("SOW Reference",   id_.sow_reference),
        ("Engagement Type", id_.engagement_type.value.replace("_", " ").title()),
        ("Authorization Basis", ", ".join(r.value for r in id_.regulatory_basis) if id_.regulatory_basis else "Internal Policy"),
    ]
    if id_.msa_reference:
        meta_rows.insert(2, ("Master Service Agreement", id_.msa_reference))
    _make_table(doc, ["Field", "Value"], [2.5, 4.5], [[k, v] for k, v in meta_rows])

    _subheading(doc, "1.3 Engagement Period")
    start = per.authorized_start_date
    end   = per.authorized_end_date
    tz    = getattr(start, 'tzinfo', None)
    tz_str = str(tz) if tz else "Eastern Time"
    days  = ", ".join(per.active_testing_days)

    period_rows = [
        ("Authorized Start Date", f"{_fmt_date(start)} — {per.active_testing_hours_start}"),
        ("Authorized End Date",   f"{_fmt_date(end)} — {per.active_testing_hours_end}"),
        ("Active Testing Window", f"{days}, {per.active_testing_hours_start} – {per.active_testing_hours_end}"),
        ("Report Draft Due",      _fmt_date(per.report_draft_due)),
        ("Final Report Due",      _fmt_date(per.report_final_due)),
    ]
    if per.blackout_dates:
        period_rows.append(("Blackout Dates", "; ".join(f"{bd.date} ({bd.reason})" for bd in per.blackout_dates)))
    if per.retest_included and per.retest_window_start:
        period_rows.append(("Retest Window", f"{_fmt_date(per.retest_window_start)} – {_fmt_date(per.retest_window_end)}"))
    _make_table(doc, ["Period", "Value"], [2.5, 4.5], period_rows)

    # ── 2. Contacts ────────────────────────────────────────────────────────────
    _heading(doc, "2. Client and Testing Team Contacts")
    _subheading(doc, "2.1 Client Authorized Representatives")

    client_contacts = [c for c in eng.contacts
                       if c.organization == id_.client_org_legal_name
                       or not any(t in c.role for t in ['engagement_lead', 'team_member'])]
    if client_contacts:
        contact_rows = []
        for c in client_contacts:
            contact_rows.append([
                c.role.replace("_", " ").title(),
                c.full_name,
                c.title or "",
                c.phone_primary or "",
                c.email or "",
            ])
        _make_table(doc,
            ["Role", "Name", "Title", "Phone", "Email"],
            [1.8, 1.4, 1.4, 1.2, 1.7],
            contact_rows)
    else:
        _body(doc, "[No client contacts defined]", italic=True)

    halt = _contact_by_role(eng, 'emergency_halt_authority')
    if halt and halt.phone_mobile:
        _body(doc,
            f"Emergency out-of-band contact for immediate test suspension: "
            f"{halt.full_name} mobile {halt.phone_mobile}. "
            f"This number is to be used exclusively in the event that active testing "
            f"must be halted immediately and normal communication channels are unavailable.",
            bold=False)

    _subheading(doc, "2.2 Testing Team Personnel")
    tester_contacts = [c for c in eng.contacts
                       if c.role in ('engagement_lead', 'team_member')]
    if tester_contacts:
        tester_rows = []
        for c in tester_contacts:
            certs = ", ".join(c.certifications) if c.certifications else ""
            ips   = ", ".join(c.authorized_source_ips) if c.authorized_source_ips else "N/A — Physical/SE"
            tester_rows.append([c.full_name, c.role.replace("_", " ").title(), certs, ips, c.email or ""])
        _make_table(doc,
            ["Name", "Role", "Certifications", "Authorized Source IP(s)", "Contact"],
            [1.6, 1.2, 1.4, 1.8, 1.5],
            tester_rows)
    else:
        _body(doc, "[No testing team contacts defined]", italic=True)

    # ── 3. Scope Definition ────────────────────────────────────────────────────
    _heading(doc, "3. Scope Definition")
    _subheading(doc, "3.1 In-Scope Networks and Hosts")
    _body(doc,
        "The following networks, VLANs, and devices are authorized for testing. "
        "Each asset is listed with its device classification, network addressing, "
        "and delivery method. All assets are individually authorized — the testing "
        "team may not expand scope to adjacent assets without written amendment.")

    if eng.in_scope_assets:
        # Group by VLAN where possible
        asset_rows = []
        for a in eng.in_scope_assets:
            # VLAN / range
            vlan_str = ""
            vid = getattr(a, 'vlan_id', None)
            vrs = getattr(a, 'vlan_range_start', None)
            vre = getattr(a, 'vlan_range_end', None)
            if vid:
                vlan_str = str(vid)
            elif vrs is not None:
                vlan_str = str(vrs) if vrs == vre else f"{vrs}–{vre or vrs}"

            raw_dt      = getattr(a, 'device_type', None) or ''
            _ACRONYMS = {'vlan':'VLAN','ids_ips':'IDS/IPS','ngfw':'NGFW',
                         'vpn_gateway':'VPN Gateway','iot_device':'IoT Device',
                         'next_gen_firewall':'Next-Gen Firewall'}
            def _fmt_dt(d):
                if d in _ACRONYMS: return _ACRONYMS[d]
                return d.replace('_',' ').title()
            device_type = ', '.join(_fmt_dt(d) for d in raw_dt) if isinstance(raw_dt, list)                           else _fmt_dt(raw_dt)
            mac         = getattr(a, 'mac_address', '') or ''
            hostname    = getattr(a, 'hostname', '') or ''
            ip_addr     = getattr(a, 'ip_address', '') or ''
            delivery    = a.delivery_method.value if a.delivery_method else ''
            os_plat  = getattr(a, 'os_platform', '') or ''
            asset_rows.append([
                a.asset_name,
                device_type,
                (a.cidr_notation or '') + (" / " + a.subnet_mask if a.subnet_mask else ""),
                ip_addr or hostname or '',
                vlan_str,
                mac,
                os_plat,
                delivery.replace('_', ' '),
                a.description or "",
            ])
        _make_table(doc,
            ["Asset / Segment", "Type", "CIDR / Mask", "IP / Host", "VLAN", "MAC", "OS / Platform", "Delivery", "Description"],
            [1.3, 0.85, 1.05, 1.0, 0.45, 0.9, 0.95, 0.65, 1.3],
            asset_rows)

        # Undisclosed device disclaimer
        _body(doc,
            "UNDISCLOSED DEVICE DISCLAIMER: The testing firm is not responsible for "
            "connectivity, discovery, or functionality of any devices not disclosed prior "
            "to test initiation. Best-effort discovery will be performed; however, "
            "undocumented devices that are left in a disconnected or dysfunctional state "
            "as a result of testing are the sole responsibility of the client to document "
            "and repair upon presentation of the testing audit trail.",
            bold=True)
    else:
        _body(doc, "[No in-scope assets defined]", italic=True)

    # Note about TBD assets
    tbd = [a for a in eng.in_scope_assets if not a.delivery_confirmed]
    if tbd:
        _body(doc,
            f"NOTE: {len(tbd)} asset(s) are client-provisioned and pending confirmed "
            f"IP address/hostname delivery. Testing of these specific assets is BLOCKED "
            f"until written delivery confirmation is received from the client.",
            bold=True)

    _subheading(doc, "3.2 Explicitly Out-of-Scope Assets")
    _body(doc,
        "The following assets are explicitly EXCLUDED from testing scope. "
        "No testing activity — including passive scanning, enumeration, or traffic "
        "generation — may target these assets. For third-party operated assets, "
        "emergency contact information is provided for immediate use in the event "
        "of accidental contact.")
    if eng.out_of_scope_assets:
        oos_rows = []
        for a in eng.out_of_scope_assets:
            vlan_str = str(getattr(a, 'vlan_id', '') or '')
            ip_str   = getattr(a, 'ip_address', '') or getattr(a, 'cidr_notation', '') or ''
            mac_str  = getattr(a, 'mac_address', '') or ''
            tp_op_name  = getattr(a, 'third_party_name', '') or ''
            tp_cname    = getattr(a, 'third_party_contact_name', '') or ''
            tp_phone    = getattr(a, 'third_party_contact_phone', '') or ''
            # Show operator name on first line, emergency contact on second
            tp_contact  = tp_op_name
            if tp_cname or tp_phone:
                tp_contact += ("\n" if tp_op_name else "") + tp_cname
                if tp_phone:
                    tp_contact += (" " if tp_cname else "") + tp_phone
            oos_rows.append([
                a.asset_name,
                ip_str,
                vlan_str,
                mac_str,
                getattr(a, 'exclusion_reason', '') or '',
                tp_contact,
            ])
        _make_table(doc,
            ["Asset", "IP / CIDR", "VLAN", "MAC", "Exclusion Reason", "3rd Party Contact"],
            [1.4, 1.2, 0.6, 1.0, 1.9, 1.4],
            oos_rows)

        # Regulatory exclusions note
        reg_excl = [a for a in eng.out_of_scope_assets if getattr(a, 'regulatory_exclusion', False)]
        if reg_excl:
            _body(doc,
                f"NOTE: {len(reg_excl)} asset(s) are excluded due to regulatory mandate "
                f"(e.g. Federal Reserve, SWIFT, PCI QSA scope boundary). Testing these assets "
                f"may constitute a regulatory violation independent of this Agreement.",
                bold=True)
    else:
        _body(doc, "[No out-of-scope assets defined]", italic=True)

    if eng.physical_locations:
        _subheading(doc, "3.3 Physical Locations In Scope")
        _body(doc,
            "Physical testing is authorized at the following locations. "
            "Authorized activities are listed per location — all other physical "
            "testing activities are NOT authorized at that location. "
            "Pre-notification requirements are mandatory.")
        for loc in eng.physical_locations:
            # Sub-section per location
            loc_type = getattr(loc, 'location_type', '') or ''
            loc_type_str = f" ({loc_type.replace('_',' ').title()})" if loc_type else ""
            _subheading(doc, f"• {loc.location_name}{loc_type_str}")
            addr = getattr(loc, 'address_full', '') or ''
            if addr:
                _body(doc, f"Address: {addr}")
            # Activities checklist
            activities = loc.authorized_activities or []
            if activities:
                _body(doc, "Authorized Activities:", bold=True)
                for act in activities:
                    _bullet(doc, act.replace('_', ' ').replace('-', ' ').title())
            # Pre-notification
            pre_req = getattr(loc, 'pre_notification_required', False)
            pre_hrs = getattr(loc, 'pre_notification_hours', None)
            if pre_req and pre_hrs:
                _body(doc,
                    f"Pre-Notification Required: {pre_hrs} hours advance notice "
                    f"to Physical Security Manager before any testing at this location.",
                    bold=True)
            # Third-party
            if getattr(loc, 'facility_third_party', False):
                contact = getattr(loc, 'facility_security_contact', '') or ''
                _body(doc,
                    f"Third-Party Facility: This location is operated by a third party. "
                    + (f"Facility security contact: {contact}." if contact else ""),
                    bold=True)

    if eng.social_engineering:
        _subheading(doc, "3.4 Social Engineering Scope")
        se = eng.social_engineering
        _body(doc,
            "Default target scope: ALL employees unless explicitly listed in the exclusion table below. "
            "Phishing, vishing, smishing, impersonation, and USB drop activities are "
            "governed by the authorization matrix and exclusion list below.")

        se_rows = [
            ("Phishing — Email",
             "AUTHORIZED" if se.phishing_authorized else "NOT AUTHORIZED",
             (f"Departments: {', '.join(se.phishing_target_departments)}" if se.phishing_target_departments else "All departments")
             + (f"\nTarget list delivery due: {_fmt_date(se.phishing_target_list_due_date)}" if se.phishing_target_list_due_date else "")),
            ("Vishing — Phone",
             "AUTHORIZED" if se.vishing_authorized else "NOT AUTHORIZED",
             se.vishing_targets or "All employees"),
            ("Smishing — SMS",
             "AUTHORIZED" if se.smishing_authorized else "NOT AUTHORIZED",
             "All employees"),
            ("Impersonation",
             "AUTHORIZED" if se.impersonation_authorized else "NOT AUTHORIZED",
             ("Approved pretexts: " + "; ".join(se.approved_pretexts)) if se.approved_pretexts else ""),
            ("USB Drop / Baiting",
             "AUTHORIZED" if se.usb_drop_authorized else "NOT AUTHORIZED",
             (f"Payload type: {se.usb_payload_type.value}" if se.usb_payload_type else "")
             + (f"  Recovery window: {se.usb_recovery_window_hours}h" if se.usb_recovery_window_hours else "")),
        ]
        _make_table(doc,
            ["Vector", "Status", "Conditions / Notes"],
            [1.6, 1.2, 4.7],
            [[r[0], r[1], r[2]] for r in se_rows])

        # Exclusion table — who is NOT a target
        excluded = se.excluded_se_targets or []
        if excluded:
            _body(doc, "Explicitly Excluded from SE Targeting:", bold=True)
            _make_table(doc,
                ["#", "Excluded Individual / Role / Department"],
                [0.4, 7.1],
                [[str(i+1), t] for i, t in enumerate(excluded)])
        else:
            _body(doc,
                "Exclusion List: NONE — All employees are in scope for social engineering testing. "
                "No individuals, roles, or departments are excluded.",
                bold=True)

    # ── CDE Section (PCI-DSS only) ────────────────────────────────────────────────
    reg_basis = [r.value for r in (id_.regulatory_basis or [])]
    if 'PCI-DSS' in reg_basis:
        cde_decision = getattr(id_, 'cde_scope_decision', None) or ''
        cde_notes    = getattr(id_, 'cde_scope_notes', '') or ''
        _subheading(doc, "3.5 PCI-DSS Cardholder Data Environment (CDE) Scope")
        _body(doc,
            "PCI-DSS is listed in the regulatory basis for this engagement. "
            "The parties have explicitly addressed the Cardholder Data Environment "
            "scope as follows:")
        decision_labels = {
            'in_scope':            'IN SCOPE — CDE assets are included in this engagement.',
            'explicitly_excluded': 'EXPLICITLY EXCLUDED — CDE assets are excluded from this engagement.',
            'not_applicable':      'NOT APPLICABLE — No CDE assets exist within the tested environment.',
        }
        decision_text = decision_labels.get(cde_decision,
            '[CDE scope decision not defined — must be resolved before document execution]')
        _body(doc, f"CDE Scope Decision: {decision_text}", bold=True)
        if cde_notes:
            _body(doc, cde_notes)
        _body(doc,
            "This determination does not constitute a PCI DSS assessment, "
            "Report on Compliance (ROC), or Self-Assessment Questionnaire (SAQ). "
            "Penetration testing results may be used as evidence in a PCI DSS assessment "
            "but do not independently establish PCI DSS compliance.",
            italic=True)
        doc.add_paragraph()

    # ── 4. Deliverables ────────────────────────────────────────────────────────
    _heading(doc, "4. Deliverables")
    deliverable_rows = [
        ("Kickoff Meeting",         "Pre-engagement call to confirm scope, contacts, and logistics",
         _fmt_date(per.authorized_start_date), "Video conference"),
        ("Weekly Status Reports",   "Brief written summary of activities, findings, and blockers",
         "Weekly during active testing", "Encrypted email"),
        ("Critical Finding Notification", "Any Critical/High finding with active exploit potential — within 4 hours",
         "Immediate upon discovery", "Phone + encrypted email"),
        ("Draft Final Report",      "Full findings report with executive summary, technical details, evidence, and remediation",
         _fmt_date(per.report_draft_due), "Encrypted email / secure portal"),
        ("Final Report",            "Finalized report incorporating client review corrections",
         _fmt_date(per.report_final_due), "Encrypted email / secure portal"),
    ]
    if per.retest_included:
        deliverable_rows.append(
            ("Retest Report", "Findings verification for remediated items only",
             _fmt_date(per.retest_window_end), "Encrypted email / secure portal"))
    _make_table(doc,
        ["Deliverable", "Description", "Due Date", "Delivery Method"],
        [1.6, 3.2, 1.2, 1.5],
        [[r[0], r[1], r[2], r[3]] for r in deliverable_rows])

    # ── 5. Data Governance ─────────────────────────────────────────────────────
    if dg:
        _heading(doc, "5. Data Governance")
        _subheading(doc, "5.1 Data Captured During Testing")
        _body(doc,
            "The Testing Team acknowledges that testing activities may result in the "
            "incidental capture of sensitive data including customer personally identifiable "
            "information (PII), account numbers, authentication credentials, and confidential "
            "business data. The Testing Team agrees to the following data handling obligations:")
        obligations = [
            f"All data captured during testing will be stored on Testing Team-controlled systems using {dg.evidence_encryption_standard.value} encryption at rest.",
            "No Client data will be stored on personal devices, cloud storage services, or any system outside the Testing Team's direct control.",
            f"Credentials discovered during testing will be reported to the Client Technical Contact within {dg.credential_reporting_window_hours} hour(s) of discovery.",
            "Customer PII encountered incidentally will be documented by type and volume only — individual records will not be retained.",
            f"All captured data will be transferred to Client and securely deleted from all Testing Team systems within {dg.evidence_retention_days} days of final report delivery.",
        ]
        if dg.pii_handling_policy:
            obligations.append(dg.pii_handling_policy)
        for o in obligations:
            _bullet(doc, o)

    # ── Framework appendix in SOW ──────────────────────────────────────────────
    # (full appendix is in ROE; SOW references with abbreviated table)
    next_section = 6 if dg else 5

    _heading(doc, f"{next_section}. Regulatory and Legal Considerations")
    _body(doc,
        "This engagement is conducted pursuant to applicable federal and state law. "
        "The following provisions define the legal framework under which testing is authorized, "
        "the protections afforded to Testing Team personnel, and the obligations of both parties.")
    _body(doc,
        "NOTICE: The following clauses include protective language specific to authorized "
        "penetration testing engagements. Both parties are advised to have this Agreement "
        "reviewed by qualified Attorney of Record before execution. These provisions are designed "
        "to protect both the Client and the Testing Team.",
        bold=True)

    legal_start = next_section + 1
    _render_clauses(doc, get_sow_legal_clauses(), start_num=legal_start)

    sig_num = legal_start + len(get_sow_legal_clauses())

    # ── Signatures ──────────────────────────────────────────────────────────────
    _heading(doc, f"{sig_num}. Signatures and Authorization")
    _body(doc,
        "By signing below, each party represents that they have read, understood, and "
        "agree to the terms of this Agreement. The Client signatory represents and warrants "
        "that they have legal authority to authorize penetration testing of all assets listed "
        "in Section 3.",
        bold=False)
    _body(doc,
        "THIS DOCUMENT MUST BE FULLY EXECUTED BY BOTH PARTIES BEFORE ANY TESTING ACTIVITY COMMENCES.",
        bold=True)
    doc.add_paragraph()

    _body(doc, f"For {id_.client_org_legal_name}", bold=True)
    _signature_table(doc,
        "Signature — Authorizing Executive (CISO)",
        "Signature — Reviewed by General Counsel")
    doc.add_paragraph()

    _body(doc, f"For {id_.testing_firm_legal_name}", bold=True)
    _signature_table(doc,
        "Signature — Engagement Lead",
        "Signature — Testing Firm Principal / Director")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# ROE GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def generate_roe(eng: Engagement) -> bytes:
    """Generate the Rules of Engagement document. Returns bytes."""
    doc = DocxDocument()

    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    id_  = eng.identity
    per  = eng.period

    # ── Headers / Footers ──────────────────────────────────────────────────────
    _add_header_footer(
        doc.sections[0],
        classification=id_.classification.value,
        doc_title="RULES OF ENGAGEMENT",
        eng_id=id_.engagement_id,
        version=id_.document_version,
    )

    # ── Cover ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RULES OF ENGAGEMENT")
    r.bold = True; r.font.size = Pt(20); r.font.name = 'Arial'
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PENETRATION TEST OPERATIONAL PROCEDURES AND CONSTRAINTS")
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'
    r.font.color.rgb = MID

    doc.add_paragraph()

    _cover_table(doc, [
        ("Companion Document",  id_.engagement_id + " (Scope of Work)"),
        ("ROE Document ID",     id_.engagement_id + "-ROE"),
        ("Client",              id_.client_org_legal_name),
        ("Testing Firm",        id_.testing_firm_legal_name),
        ("Effective Date",      _fmt_date(per.authorized_start_date)),
        ("Expiration Date",     _fmt_date(per.retest_window_end or per.authorized_end_date)),
        ("Classification",      f"{id_.classification.value.upper()} — RESTRICTED"),
    ])

    _notice_box(doc,
        "This Rules of Engagement document defines the operational constraints, "
        "communication procedures, escalation requirements, and technique-level "
        "authorization matrix for this engagement. This document must be read in "
        "conjunction with the Scope of Work. In the event of conflict, the more "
        "restrictive provision governs.")

    # ── 1. Communication Protocols ─────────────────────────────────────────────
    _heading(doc, "1. Communication Protocols")
    _subheading(doc, "1.1 Daily Operational Communication")

    lead = _contact_by_role(eng, 'engagement_lead')
    tech = _contact_by_role(eng, 'primary_technical_contact')
    _body(doc,
        f"The {'Engagement Lead' if not lead else lead.full_name} will send a brief daily "
        f"activity summary to the Client Technical Contact "
        f"({'[TBD]' if not tech else tech.full_name}) by 18:00 Eastern each active testing day. "
        f"This summary will include: systems tested, techniques employed, anomalies encountered, "
        f"and planned activities for the following day.")
    _body(doc, "All daily communications will be sent via encrypted email.")

    _subheading(doc, "1.2 Critical Finding Notification (4-Hour Rule)")
    _body(doc,
        "If the Testing Team discovers any finding meeting the following criteria, "
        "the Client CISO must be notified within 4 hours of discovery regardless of time of day:")
    critical_triggers = [
        "Any unauthenticated remote code execution against a production system",
        "Any confirmed access to customer PII, account data, or transaction records",
        "Any successful authentication bypass against online banking or core systems",
        "Any finding that poses an immediate risk of financial loss or regulatory violation",
        "Any evidence of an existing third-party compromise or active intrusion",
        "Any accidental impact to production system availability or data integrity",
    ]
    for t in critical_triggers:
        _bullet(doc, t)

    ciso = _contact_by_role(eng, 'authorizing_executive')
    bcc  = _contact_by_role(eng, 'business_continuity_contact')
    if ciso:
        _body(doc,
            f"Critical finding notification will be made via phone call to the Client CISO "
            f"({ciso.full_name}, {ciso.phone_primary}) followed immediately by encrypted email."
            + (f" If the CISO is unreachable after two call attempts within 30 minutes, "
               f"{bcc.full_name} ({bcc.phone_primary}) will be contacted." if bcc else ""))

    _subheading(doc, "1.3 Emergency Test Suspension")
    _body(doc, "Testing must be immediately suspended and the Client CISO notified if any of the following occur:")
    suspension_triggers = [
        "Client invokes the stop-work order via the emergency contact number",
        "Testing Team detects that testing activity has caused unintended production impact",
        "Testing Team gains access to an out-of-scope system through lateral movement or misconfiguration",
        "Testing Team discovers evidence of an active breach by a third party",
        "Any member of the testing team is challenged, detained, or threatened during physical testing",
    ]
    for t in suspension_triggers:
        _bullet(doc, t)
    _body(doc,
        "Following suspension, the Testing Team will provide a written incident report within "
        "2 business hours. No testing may resume without explicit written authorization from "
        "the Client CISO.")

    # ── 2. Technique Authorization Matrix ─────────────────────────────────────
    _heading(doc, "2. Technique Authorization Matrix")
    _body(doc,
        "The following matrix defines authorization status for each testing technique category. "
        "AUTHORIZED = authorized as described. NOT AUTHORIZED = strictly prohibited. "
        "CONDITIONAL = authorized only under the stated conditions. "
        "Any technique not listed is considered NOT AUTHORIZED by default.")

    categories = ['reconnaissance', 'vuln_scanning', 'exploitation',
                  'post_exploitation', 'dos', 'social_engineering', 'physical']
    cat_labels  = {
        'reconnaissance':    '2.1 Reconnaissance',
        'vuln_scanning':     '2.2 Vulnerability Scanning',
        'exploitation':      '2.3 Exploitation',
        'post_exploitation': '2.4 Post-Exploitation',
        'dos':               '2.5 Denial of Service',
        'social_engineering':'2.6 Social Engineering',
        'physical':          '2.7 Physical',
    }

    _body(doc,
        "Technique IDs reference MITRE ATT\u2019\u00e9 ATT&CK Enterprise v15. "
        "Testing phases follow NIST SP 800-115 and PTES. "
        "See Appendix A for framework reference.")

    TACTIC_NAMES = {
        "TA0043": "Reconnaissance", "TA0042": "Resource Development",
        "TA0001": "Initial Access",  "TA0002": "Execution",
        "TA0003": "Persistence",     "TA0004": "Privilege Escalation",
        "TA0005": "Defense Evasion", "TA0006": "Credential Access",
        "TA0007": "Discovery",       "TA0008": "Lateral Movement",
        "TA0009": "Collection",      "TA0010": "Exfiltration",
        "TA0011": "Command and Control", "TA0040": "Impact",
    }

    for cat in categories:
        techs = [t for t in eng.techniques if t.category.value == cat]
        if not techs:
            continue
        _subheading(doc, cat_labels.get(cat, cat.replace("_", " ").title()))
        rows = []
        for t in techs:
            status = t.authorization_status.value.upper().replace("_", " ")
            if t.prohibited:
                status = "NOT AUTHORIZED"

            # ATT&CK reference
            att_ref = ""
            if hasattr(t, 'mitre_technique_id') and t.mitre_technique_id:
                tactic = getattr(t, 'mitre_tactic_id', '') or ''
                tactic_name = TACTIC_NAMES.get(tactic, tactic)
                att_ref = f"{t.mitre_technique_id}"
                if tactic and tactic != 'N/A':
                    att_ref += f" ({tactic_name})"

            conditions = ""
            if t.authorization_status == AuthorizationStatus.CONDITIONAL:
                conditions = t.conditions or ""
                if t.scope_limitation:
                    conditions = (conditions + " " + t.scope_limitation).strip()
            elif t.scope_limitation:
                conditions = t.scope_limitation
            if t.prohibited:
                conditions = conditions or "Strictly prohibited."

            rows.append([t.technique_name, att_ref, status, conditions])
        _make_table(doc,
            ["Technique", "ATT&CK ID", "Auth", "Conditions / Constraints"],
            [2.0, 1.3, 1.0, 3.2],
            rows)

    # ── 3. Prohibited Actions ──────────────────────────────────────────────────
    _heading(doc, "3. Prohibited Actions (Absolute)")
    _body(doc,
        "The following actions are prohibited under all circumstances regardless of "
        "technical feasibility, discovery context, or perceived justification. Violation "
        "of any prohibition below constitutes a material breach of this Agreement and may "
        "result in immediate engagement termination, legal action, and regulatory notification.")

    prohibited_techs = [t for t in eng.techniques if t.prohibited]
    standard_prohibitions = [
        "Accessing, copying, modifying, or exfiltrating actual customer financial data, PII, or account records",
        "Executing any financial transaction, transfer, or account modification on any system",
        "Encrypting, destroying, corrupting, or rendering inaccessible any Client data or system",
        "Installing persistent backdoors or remote access tools on production systems without explicit written authorization",
        "Deleting, altering, or tampering with system logs or audit trails",
        "Exploiting zero-day vulnerabilities (undisclosed vulnerabilities must be reported, not exploited)",
        "Targeting Client customers, vendors, regulators, or any individual not employed by the Client",
        "Physical property damage, forced entry, or physical confrontation with any person",
        "Continuing testing after a stop-work order has been issued",
        "Testing from any source IP not listed in the SOW without prior written approval",
        "Disclosing, publishing, or sharing any engagement findings with any third party without written Client authorization",
    ]
    for p_text in standard_prohibitions:
        _bullet(doc, p_text)
    for t in prohibited_techs:
        if t.scope_limitation:
            _bullet(doc, t.scope_limitation)

    # ── 4. Maintenance Windows ─────────────────────────────────────────────────
    if eng.maintenance_windows:
        _heading(doc, "4. Maintenance Windows")
        _body(doc,
            "The following maintenance windows are pre-authorized for disruptive testing "
            "activities including DoS simulation, load testing, and production system exploitation "
            "(where separately authorized). All times are Eastern Time.")

        # Window schedule table
        mw_rows = []
        for mw in eng.maintenance_windows:
            activities  = ", ".join(mw.authorized_activity_refs) if mw.authorized_activity_refs else ""
            ids_status  = "Whitelisted ✓" if getattr(mw, 'ids_ips_whitelisted', False) else                           ("Present — NOT whitelisted ⚠" if getattr(mw, 'ids_ips_present', False) else "Not applicable")
            soc_status  = "Notified ✓" if getattr(mw, 'soc_notified', False) else "NOT notified ⚠"
            mw_rows.append([
                mw.window_id,
                str(mw.date),
                f"{mw.start_time} – {mw.end_time}",
                activities,
                ids_status,
                soc_status,
            ])
        _make_table(doc,
            ["Window ID", "Date", "Time", "Authorized Activities", "IDS/IPS", "SOC"],
            [0.9, 1.1, 1.4, 2.3, 1.4, 1.0],
            mw_rows)

        first_mw = eng.maintenance_windows[0]
        _bullet(doc,
            f"Pre-notification: Engagement Lead must confirm intent to use a maintenance window "
            f"with the Client Technical Contact no later than "
            f"{first_mw.pre_notification_hours} hours before window start time.")
        _bullet(doc,
            f"Cancellation: Client may cancel with "
            f"{first_mw.cancellation_notice_hours}-hour notice by phone and email.")
        _bullet(doc,
            "IDS/IPS: Any IDS/IPS or WAF rules must be confirmed whitelisted for testing "
            "source IPs before each window opens. Testing that triggers incident response "
            "due to unconfigured IDS/IPS rules does not constitute unauthorized access.")
        _bullet(doc,
            "SOC Notification: The Security Operations Center duty contact must be "
            "notified before each window. Testing traffic during a notified window must "
            "not be treated as a real incident.")
        _bullet(doc,
            "Staffing: Engagement Lead and Client Technical Contact must be reachable "
            "by phone throughout any window in which disruptive testing occurs.")

        # Per-window pre-start checklists
        for mw in eng.maintenance_windows:
            checklist = getattr(mw, 'pre_window_checklist', None) or []
            if checklist:
                _body(doc, f"Pre-Start Checklist for {mw.window_id}:", bold=True)
                for item in checklist:
                    _bullet(doc, f"☐  {item}")

    # ── 5. Evidence Collection ─────────────────────────────────────────────────
    _heading(doc, "5. Evidence Collection and Handling")
    _subheading(doc, "5.1 Required Evidence Standards")
    evidence_standards = [
        "Screenshots must include timestamp, hostname or IP of the target system, and the tester's workstation identifier.",
        "Command-line output must be captured in full — truncated output is not acceptable as sole evidence.",
        "Tool output logs must be saved in their native format in addition to any screenshots.",
        "For exploitation findings: evidence must demonstrate the complete attack chain from initial access through the impact demonstrated.",
        "All evidence files must be named with the format: [EngagementID]-[FindingID]-[EvidenceType]-[Sequence].",
    ]
    for s in evidence_standards:
        _bullet(doc, s)

    if eng.data_governance:
        _subheading(doc, "5.2 Evidence Storage and Security")
        dg = eng.data_governance
        _bullet(doc, f"All evidence stored on Testing Team-controlled, encrypted systems ({dg.evidence_encryption_standard.value}) throughout the engagement.")
        _bullet(doc, "No evidence stored on personal devices, cloud consumer storage, or any Client-operated systems.")
        _bullet(doc, f"Evidence transferred to Client via secure channel upon final report delivery.")
        _bullet(doc, f"Testing Team evidence copies securely deleted within {dg.evidence_retention_days} days of final report delivery. Written deletion confirmation provided to Client.")

    # ── Framework Reference Appendix ───────────────────────────────────────────
    roe_appendix_num = 6 if not eng.maintenance_windows else 7
    _heading(doc, f"Appendix A. Framework References")
    _body(doc,
        "This document was produced in alignment with the following industry-standard "
        "frameworks and references. Use of these frameworks ensures the engagement "
        "follows recognized, defensible methodology.")

    framework_rows = [
        ("MITRE ATT&CK Enterprise v15",
         "https://attack.mitre.org",
         "Technique taxonomy and tactic/technique IDs used throughout the authorization matrix. "
         "ATT&CK is maintained by MITRE Corporation and is the de facto standard for adversary "
         "behavior classification."),
        ("NIST SP 800-115",
         "https://csrc.nist.gov/publications/detail/sp/800-115/final",
         "Technical Guide to Information Security Testing and Assessment. Published by the "
         "National Institute of Standards and Technology. Defines testing phases (Planning, "
         "Discovery, Attack, Reporting) referenced in the technique matrix."),
        ("PTES — Penetration Testing Execution Standard",
         "http://www.pentest-standard.org",
         "Community standard defining pre-engagement, intelligence gathering, threat modeling, "
         "vulnerability analysis, exploitation, post-exploitation, and reporting phases."),
        ("OWASP Testing Guide v4.2",
         "https://owasp.org/www-project-web-security-testing-guide/",
         "Open Web Application Security Project testing methodology. Referenced for web "
         "application and API testing techniques."),
        ("CVSS v3.1",
         "https://www.first.org/cvss/",
         "Common Vulnerability Scoring System. Findings in the final report will be scored "
         "using CVSS v3.1 Base Scores. Published by FIRST (Forum of Incident Response and "
         "Security Teams)."),
        ("CVE / NVD",
         "https://nvd.nist.gov",
         "Known vulnerabilities referenced in findings will be cited by CVE identifier where "
         "applicable. The National Vulnerability Database is maintained by NIST."),
        ("CWE",
         "https://cwe.mitre.org",
         "Common Weakness Enumeration. Root-cause weakness classifications for findings will "
         "reference CWE IDs where applicable."),
    ]
    _make_table(doc,
        ["Framework", "Reference", "Applicability"],
        [2.0, 2.0, 3.5],
        framework_rows)

    # ── Legal sections ─────────────────────────────────────────────────────────
    roe_legal_start = roe_appendix_num + 1
    _render_clauses(doc, get_roe_legal_clauses(), start_num=roe_legal_start)

    roe_sig_num = roe_legal_start + len(get_roe_legal_clauses())

    # ── Signatures ──────────────────────────────────────────────────────────────
    _heading(doc, f"{roe_sig_num}. Signatures")
    _body(doc,
        "By signing below, both parties acknowledge that they have read, understood, and "
        "agree to operate within the constraints defined in this Rules of Engagement document "
        "for the duration of the engagement.")
    _body(doc,
        "THESE RULES ARE BINDING ON ALL TESTING TEAM PERSONNEL. THE ENGAGEMENT LEAD IS "
        "RESPONSIBLE FOR ENSURING ALL TEAM MEMBERS HAVE READ AND UNDERSTOOD THIS DOCUMENT "
        "BEFORE CONDUCTING ANY TESTING ACTIVITY.",
        bold=True)
    doc.add_paragraph()

    _body(doc, f"For {id_.client_org_legal_name}", bold=True)
    _signature_table(doc,
        "Signature — CISO (Authorizing Executive)",
        "Signature — Primary Technical Contact")
    doc.add_paragraph()

    _body(doc, f"For {id_.testing_firm_legal_name}", bold=True)
    _signature_table(doc,
        "Signature — Engagement Lead",
        "Signature — Testing Firm Principal / Director")

    # Team acknowledgment table
    testers = [c for c in eng.contacts if c.role in ('engagement_lead', 'team_member')]
    if testers:
        doc.add_paragraph()
        _body(doc,
            "Each testing team member listed below must initial to confirm they have read "
            "and agree to these Rules of Engagement:",
            bold=False)
        ack_rows = [[c.full_name, c.role.replace("_", " ").title(), "", ""] for c in testers]
        _make_table(doc,
            ["Name", "Role", "Initial", "Date"],
            [2.5, 2.5, 1.5, 1.0],
            ack_rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
